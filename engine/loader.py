"""
文档加载器 — 支持 ZIP / PDF / Word 解析

PDF 处理:
- 优先提取电子文本 (PyMuPDF)
- 无文字时自动 OCR 识别扫描件 (PaddleOCR, 中文场景准确率 ~98%)
- 表格单独识别

输入: 文件路径 (支持 .zip / .pdf / .docx)
输出: 标准化 Document 对象列表
"""

import io
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import numpy as np

import config


@dataclass
class Document:
    """标准化的文档片段"""
    content: str                          # 文本内容
    source: str                           # 来源文件名
    page: int = 1                         # 页码 (PDF 有效)
    heading: str = ""                     # 所属章节标题
    is_table: bool = False                # 是否表格内容
    metadata: dict = field(default_factory=dict)


def load_file(file_path: str | Path) -> list[Document]:
    """
    根据文件扩展名自动选择解析器
    支持: .pdf / .docx / .zip (内含 pdf/docx)
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _load_pdf(path)
    elif ext in (".docx", ".doc"):
        return _load_docx(path)
    elif ext == ".zip":
        return _load_zip(path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，请上传 PDF / Word / ZIP")


def _load_pdf(path: Path) -> list[Document]:
    """
    解析 PDF，按页提取文本
    - 提取字体/字号/位置信息用于标题识别 (get_text("dict"))
    - 优先电子文本，无文字时自动 OCR (PaddleOCR)
    - 表格单独识别
    """
    docs = []
    doc = fitz.open(path)
    source = path.name
    ocr_pages = 0
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text").strip()

            # --- 电子文本为空 → OCR 扫描件 ---
            is_ocr = False
            if not text:
                text = _ocr_page(page)
                if text:
                    is_ocr = True
                    ocr_pages += 1

            if not text:
                continue

            # --- 提取字体信息用于标题识别 ---
            font_info = _extract_font_info(page) if not is_ocr else []

            # 尝试提取表格
            tables = page.find_tables()
            table_texts = _extract_table_texts(tables)

            if table_texts:
                for t_text in table_texts:
                    docs.append(Document(
                        content=t_text,
                        source=source,
                        page=page_num + 1,
                        is_table=True,
                        metadata={"type": "table", "ocr": is_ocr}
                    ))

            if text:
                docs.append(Document(
                    content=text,
                    source=source,
                    page=page_num + 1,
                    metadata={"type": "text", "ocr": is_ocr, "font_info": font_info}
                ))

        if ocr_pages > 0:
            print(f"[PDF] {source}: {ocr_pages}/{len(doc)} 页使用 OCR 识别")
    finally:
        doc.close()
    return docs


def _extract_font_info(page: fitz.Page) -> list[dict]:
    """
    从 PDF 页面提取每行文字的字体信息
    用于后续智能识别标题层级

    Returns: [{"text": "第一章", "size": 16.0, "bold": True, "y": 100.5}, ...]
    """
    blocks = page.get_text("dict")["blocks"]
    lines_info = []

    for block in blocks:
        if block.get("type") != 0:  # 非文本块跳过
            continue
        for line in block.get("lines", []):
            spans = line["spans"]
            if not spans:
                continue

            # 合并同一行的所有 span
            line_text = "".join(s["text"] for s in spans)
            # 取最大字号和是否粗体
            max_size = max(s["size"] for s in spans)
            is_bold = any(s.get("flags", 0) & 2 for s in spans)  # flags bit 1 = bold
            y_pos = line["bbox"][1]  # 顶部 y 坐标

            if line_text.strip():
                lines_info.append({
                    "text": line_text.strip(),
                    "size": round(max_size, 1),
                    "bold": is_bold,
                    "y": round(y_pos, 1),
                })

    return lines_info


def _ocr_page(page: fitz.Page) -> str:
    """
    将 PDF 页渲染为图像，用 PaddleOCR 识别文字
    中文场景准确率远优于 Tesseract (~98% vs ~85%)
    """
    try:
        from paddleocr import PaddleOCR
    except ImportError:
        return ""   # 未安装 paddleocr，跳过

    try:
        # PyMuPDF 渲染为 numpy array
        pix = page.get_pixmap(dpi=config.OCR_DPI)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, pix.n
        )

        # 单例 PaddleOCR，避免每页重新初始化
        ocr = _get_paddle_ocr()

        result = ocr.ocr(img)
        if not result or not result[0]:
            return ""

        # 拼接所有识别行
        lines = [line[1][0] for line in result[0] if line[1][1] > 0.5]  # 置信度 > 0.5
        return "\n".join(lines).strip()
    except Exception:
        return ""


_paddle_ocr_instance = None


def _get_paddle_ocr():
    """PaddleOCR 单例，首次初始化约 1-2 秒"""
    global _paddle_ocr_instance
    if _paddle_ocr_instance is None:
        from paddleocr import PaddleOCR
        _paddle_ocr_instance = PaddleOCR(lang=config.OCR_LANG)
    return _paddle_ocr_instance


def _load_docx(path: Path) -> list[Document]:
    """解析 Word 文档，按段落提取"""
    docs = []
    source = path.name
    doc = DocxDocument(str(path))

    # 提取段落
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        docs.append(Document(
            content="\n".join(paragraphs),
            source=source,
            metadata={"type": "text"}
        ))

    # 提取表格
    for idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            docs.append(Document(
                content=table_text,
                source=source,
                is_table=True,
                metadata={"type": "table", "table_index": idx}
            ))

    return docs


def _load_zip(path: Path) -> list[Document]:
    """解压 ZIP 并递归解析其中的 PDF/Word 文件"""
    all_docs = []
    upload_dir = config.UPLOAD_DIR

    with zipfile.ZipFile(path, "r") as zf:
        # 解压到 uploads 目录
        extract_dir = upload_dir / path.stem
        zf.extractall(extract_dir)

        # 遍历解压后的文件
        for extracted in extract_dir.rglob("*"):
            if extracted.is_file():
                try:
                    docs = load_file(extracted)
                    all_docs.extend(docs)
                except ValueError:
                    # 跳过不支持的文件
                    continue

    return all_docs


def _extract_table_texts(tables) -> list[str]:
    """将 PyMuPDF 表格对象转为可读文本"""
    result = []
    for table in tables:
        try:
            data = table.extract()
            if not data:
                continue
            rows = []
            for row in data:
                # 过滤掉 None 单元格
                cells = [str(c) if c is not None else "" for c in row]
                rows.append(" | ".join(cells))
            text = "\n".join(rows)
            if text.strip():
                result.append(text)
        except Exception:
            continue
    return result


def stream_documents(file_path: str | Path) -> Iterator[Document]:
    """流式 yield Document，适合大文件处理时显示进度"""
    for doc in load_file(file_path):
        yield doc


def scan_folder(folder_path: str | Path) -> list[Path]:
    """
    递归扫描文件夹，返回所有 PDF/Word 文件路径
    用于本地大文件批量导入（不走浏览器上传）
    """
    folder = Path(folder_path)
    if not folder.exists():
        raise FileNotFoundError(f"文件夹不存在: {folder_path}")
    if not folder.is_dir():
        raise NotADirectoryError(f"不是文件夹: {folder_path}")

    supported = {".pdf", ".docx", ".doc"}
    files = []
    for f in folder.rglob("*"):
        if f.is_file() and f.suffix.lower() in supported:
            files.append(f)
    return sorted(files)


def load_folder(folder_path: str | Path) -> Iterator[Document]:
    """
    流式加载整个文件夹的文档
    逐个文件 yield Document，适合大文件夹处理时显示进度
    """
    files = scan_folder(folder_path)
    for file_path in files:
        try:
            for doc in load_file(file_path):
                yield doc
        except Exception:
            # 跳过解析失败的文件，继续处理后续文件
            continue
